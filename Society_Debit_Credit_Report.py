import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import pyodbc
import pandas as pd
from docx import Document
import tempfile
import webbrowser


class MonthDebitCreditReport:

    def __init__(self, parent):

        self.win = tk.Toplevel(parent)
        self.win.title("Month Wise Debit Credit Report")
        self.win.geometry("950x520")

        # DATABASE CONNECTION
        self.conn = pyodbc.connect(
            "DRIVER={SQL Server};"
            "SERVER=Sarthak;"
            "DATABASE=Socity_db;"
            "Trusted_Connection=yes;"
        )

        self.df = pd.DataFrame()

        self.create_widgets()

    # -----------------------------------
    # GUI
    # -----------------------------------
    def create_widgets(self):

        style = ttk.Style()
        style.configure("Treeview.Heading", font=("Arial", 11, "bold"))

        top = tk.Frame(self.win)
        top.pack(pady=10)

        tk.Label(top, text="Year").grid(row=0, column=0)

        self.year_var = tk.IntVar(value=2026)

        ttk.Combobox(
            top,
            textvariable=self.year_var,
            values=(2024, 2025, 2026, 2027),
            width=8
        ).grid(row=0, column=1, padx=5)

        tk.Label(top, text="Month").grid(row=0, column=2)

        self.month_var = tk.StringVar()

        months = [
            "January","February","March","April","May","June",
            "July","August","September","October","November","December"
        ]

        ttk.Combobox(
            top,
            textvariable=self.month_var,
            values=months,
            width=12
        ).grid(row=0, column=3, padx=5)

        tk.Label(top, text="Flat Search").grid(row=0, column=4)

        self.search_var = tk.StringVar()

        tk.Entry(top, textvariable=self.search_var, width=12).grid(row=0, column=5)

        tk.Button(top, text="Search", command=self.search_flat)\
            .grid(row=0, column=6, padx=5)

        tk.Button(
            top,
            text="Load Report",
            command=self.load_report,
            bg="#3498db",
            fg="white"
        ).grid(row=0, column=7, padx=10)

        # TABLE
        columns = ("Flat", "Debit", "Member", "Credit", "Balance")

        self.tree = ttk.Treeview(self.win, columns=columns, show="headings")

        for col in columns:
            self.tree.heading(col, text=col)

        self.tree.column("Flat", width=80, anchor="center")
        self.tree.column("Debit", width=150, anchor="e")
        self.tree.column("Member", width=350, anchor="w")
        self.tree.column("Credit", width=150, anchor="e")
        self.tree.column("Balance", width=150, anchor="e")

        self.tree.pack(fill="both", expand=True, padx=10, pady=10)

        self.tree.tag_configure("total", font=("Arial", 11, "bold"))
        self.tree.tag_configure("balance", font=("Arial", 12, "bold"))

        # BUTTONS
        bottom = tk.Frame(self.win)
        bottom.pack(pady=10)

        tk.Button(bottom, text="Export Excel",
                  command=self.export_excel,
                  bg="#27ae60",
                  fg="white",
                  width=15).grid(row=0, column=0, padx=10)

        tk.Button(bottom, text="Export Word",
                  command=self.export_word,
                  bg="#8e44ad",
                  fg="white",
                  width=15).grid(row=0, column=1, padx=10)

        tk.Button(bottom, text="Print Ledger",
                  command=self.print_ledger,
                  bg="#e67e22",
                  fg="white",
                  width=15).grid(row=0, column=2, padx=10)

    # -----------------------------------
    # LOAD REPORT
    # -----------------------------------
    def load_report(self):

        year = self.year_var.get()
        month = self.month_var.get()

        if month == "":
            messagebox.showwarning("Select Month", "Please select a month")
            return

        month_num = pd.to_datetime(month, format="%B").month

        query = f"""
        SELECT 
            m.Flat_No,
            m.Member_Name,
            MONTH(d.Date1) AS MonthNo,
            ISNULL(SUM(d.Debit),0) AS Debit,
            ISNULL(SUM(d.Credit),0) AS Credit

        FROM Member_Mast m
        LEFT JOIN Dailytrans1 d
            ON m.Flat_No = d.Flat_No
            AND YEAR(d.Date1) = {year}
            AND MONTH(d.Date1) <= {month_num}

        GROUP BY m.Flat_No, m.Member_Name, MONTH(d.Date1)
        ORDER BY m.Flat_No, MonthNo
        """

        df = pd.read_sql(query, self.conn)

        if df.empty:
            messagebox.showinfo("No Data", "No transactions found")
            return

        final_data = []

        for flat, group in df.groupby("Flat_No"):

            group = group.sort_values("MonthNo")

            balance = 0

            for _, row in group.iterrows():
                balance += row["Credit"] - row["Debit"]

            final_data.append({
                "Flat": flat,
                "Member": group.iloc[-1]["Member_Name"],
                "Debit": group["Debit"].sum(),
                "Credit": group["Credit"].sum(),
                "Balance": balance
            })

        final_df = pd.DataFrame(final_data)

        self.df = final_df
        self.show_table(final_df)

    # -----------------------------------
    # SHOW TABLE
    # -----------------------------------
    def show_table(self, data):

        for row in self.tree.get_children():
            self.tree.delete(row)

        if data.empty:
            return

        total_debit = data["Debit"].sum()
        total_credit = data["Credit"].sum()
        total_balance = data["Balance"].sum()

        for _, r in data.iterrows():
            self.tree.insert("", tk.END, values=(
                r["Flat"],
                f"{r['Debit']:,.2f}",
                r["Member"],
                f"{r['Credit']:,.2f}",
                f"{r['Balance']:,.2f}"
            ))

        self.tree.insert("", tk.END, values=("", "", "", "", ""))

        self.tree.insert("", tk.END, values=(
            "TOTAL",
            f"{total_debit:,.2f}",
            "",
            f"{total_credit:,.2f}",
            f"{total_balance:,.2f}"
        ), tags=("total",))

    # -----------------------------------
    # SEARCH
    # -----------------------------------
    def search_flat(self):

        text = self.search_var.get().strip()

        if text == "":
            data = self.df
        else:
            data = self.df[self.df["Flat"].astype(str).str.contains(text)]

        self.show_table(data)

    # -----------------------------------
    # EXPORT EXCEL
    # -----------------------------------
    def export_excel(self):

        if self.df.empty:
            messagebox.showwarning("No Data", "Load report first")
            return

        file = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel File", "*.xlsx")]
        )

        if not file:
            return

        self.df.to_excel(file, index=False)

        messagebox.showinfo("Export", "Excel report exported successfully")

    # -----------------------------------
    # EXPORT WORD
    # -----------------------------------
    def export_word(self):

        if self.df.empty:
            messagebox.showwarning("No Data", "Load report first")
            return

        file = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")]
        )

        if not file:
            return

        document = Document()
        document.add_heading("Month Wise Debit Credit Report", 1)

        table = document.add_table(rows=1, cols=5)
        headers = ["Flat", "Debit", "Member", "Credit", "Balance"]

        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for _, row in self.df.iterrows():
            r = table.add_row().cells
            r[0].text = str(row["Flat"])
            r[1].text = f"{row['Debit']:,.2f}"
            r[2].text = str(row["Member"])
            r[3].text = f"{row['Credit']:,.2f}"
            r[4].text = f"{row['Balance']:,.2f}"

        document.save(file)

        messagebox.showinfo("Export", "Word exported successfully")

    # -----------------------------------
    # PRINT LEDGER
    # -----------------------------------
    def print_ledger(self):

        html = "<h2>Month Wise Debit Credit Ledger</h2>"
        html += self.df.to_html(index=False)

        file = tempfile.NamedTemporaryFile(delete=False, suffix=".html")

        with open(file.name, "w", encoding="utf-8") as f:
            f.write(html)

        webbrowser.open(file.name)


# -----------------------------------
# MAIN PROGRAM
# -----------------------------------
if __name__ == "__main__":

    root = tk.Tk()
    root.title("Society Management System")
    root.geometry("300x200")

    tk.Button(
        root,
        text="Open Debit Credit Report",
        command=lambda: MonthDebitCreditReport(root),
        font=("Arial", 12),
        width=24,
        bg="#2ecc71",
        fg="white"
    ).pack(pady=60)

    root.mainloop()